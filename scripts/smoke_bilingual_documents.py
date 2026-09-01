"""Smoke test hors ligne des quatre aperçus importés du kit bilingue."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import text

from dashboard.rocky.config import Settings
from dashboard.rocky.database import create_db_engine, initialize_database
from dashboard.rocky.profile_documents import (
    ROCKY_MARKER,
    save_uploaded_profile_document,
    validate_pdf,
)
from dashboard.rocky.repository import RockyRepository


PROJECT_DIR = Path(__file__).resolve().parents[1]


def _create_cv(path: Path, title: str, heading: str, description: str) -> None:
    """Produit un petit CV PDF de fixture dans la langue demandée."""
    canvas = Canvas(str(path), pagesize=A4)
    canvas.drawString(72, 780, title)
    canvas.drawString(72, 750, heading)
    canvas.drawString(72, 730, description)
    canvas.save()


def _create_letter(path: Path, title: str, introduction: str, conclusion: str) -> None:
    """Produit un modèle DOCX importé avec son emplacement Rocky intact."""
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(introduction)
    document.add_paragraph(ROCKY_MARKER)
    document.add_paragraph(conclusion)
    document.save(path)


def main() -> int:
    """Importe les kits FR/EN et valide leurs quatre aperçus PDF courants."""
    with tempfile.TemporaryDirectory(prefix="rocky_bilingual_smoke_") as folder:
        root = Path(folder)
        settings = Settings(
            project_dir=PROJECT_DIR,
            database_url_override=f"sqlite:///{root / 'smoke.db'}",
            storage_dir_override=str(root / "data"),
        )
        engine = create_db_engine(settings)
        initialize_database(engine, settings)
        with engine.begin() as connection:
            user_id = int(
                connection.execute(
                    text(
                        "INSERT INTO users (email, status, email_verified_at) "
                        "VALUES ('smoke@rocky.test', 'ACTIVE', CURRENT_TIMESTAMP) "
                        "RETURNING id"
                    )
                ).scalar_one()
            )
        repository = RockyRepository(engine).for_user(user_id)
        profile_id = repository.create_profile(
            "Smoke bilingue", onboarding_status="DRAFT"
        )
        cv_path = root / "cv_fr.pdf"
        letter_path = root / "letter_fr.docx"
        english_cv_path = root / "cv_en.pdf"
        english_letter_path = root / "letter_en.docx"
        _create_cv(cv_path, "Profil Data", "EXPÉRIENCE", "Analyse de données")
        _create_letter(
            letter_path,
            "Lettre de motivation",
            "Introduction factuelle.",
            "Conclusion.",
        )
        _create_cv(english_cv_path, "Data profile", "EXPERIENCE", "Data analysis")
        _create_letter(
            english_letter_path, "Cover letter", "Factual introduction.", "Conclusion."
        )
        save_uploaded_profile_document(
            settings, repository, user_id, profile_id, "fr", "cv", cv_path.read_bytes()
        )
        save_uploaded_profile_document(
            settings,
            repository,
            user_id,
            profile_id,
            "fr",
            "letter",
            letter_path.read_bytes(),
        )
        save_uploaded_profile_document(
            settings,
            repository,
            user_id,
            profile_id,
            "en",
            "cv",
            english_cv_path.read_bytes(),
        )
        save_uploaded_profile_document(
            settings,
            repository,
            user_id,
            profile_id,
            "en",
            "letter",
            english_letter_path.read_bytes(),
        )
        documents = repository.fetch_profile_documents(profile_id)
        if len(documents) != 4:
            raise RuntimeError(
                "Le smoke test attend exactement quatre documents courants."
            )
        for document in documents:
            preview = Path(document.preview_pdf_path or document.source_path)
            if not preview.is_absolute():
                preview = settings.project_dir / preview
            validate_pdf(preview)
        print("OK · quatre aperçus PDF FR/EN importés et validés avec LibreOffice")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
