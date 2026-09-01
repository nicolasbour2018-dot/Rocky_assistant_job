"""Documents bilingues du profil, de l'import à l'aperçu PDF vérifié."""

from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
from collections.abc import Iterable
from html import escape
from pathlib import Path

from pypdf import PdfReader

from .ats import extract_pdf_text, repair_spaced_pdf_text
from .config import Settings
from .errors import DocumentError, RockyError
from .llm import RockyLLM
from .models import DocumentKind, ProfileAnalysis
from .profile_skills import infer_profile_skills_from_cv
from .repository import RockyRepository
from .text_utils import project_relative

ROCKY_MARKER = "[paragraphe Rocky]"
PROTECTED_MARKER = "__ROCKY_PARAGRAPH__"
# Des lots courts évitent qu'un CV dense dépasse la taille de sortie JSON du
# modèle et perde des lignes lors de la traduction structurée.
CV_TRANSLATION_BATCH_SIZE = 10


def _load_docx(path: Path):
    """Importe python-docx à la demande pour garder les diagnostics légers."""
    try:
        from docx import Document
    except ImportError as error:
        raise DocumentError(
            "Installe python-docx pour lire les lettres DOCX."
        ) from error
    return Document(str(path))


def file_sha256(path: Path) -> str:
    """Calcule une empreinte stable sans exposer le contenu du document."""
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _iter_cell_blocks(cell) -> Iterable:
    """Parcourt récursivement les cellules DOCX pour préserver tous les blocs de lettre."""
    yield from cell.paragraphs
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from _iter_cell_blocks(nested_cell)


def iter_docx_paragraphs(document) -> Iterable:
    """Parcourt corps, tableaux, en-têtes et pieds sans perdre un marqueur morcelé."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_cell_blocks(cell)
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def validate_pdf(path: Path) -> None:
    """Refuse les faux PDF, les documents chiffrés et les fichiers sans page."""
    if not path.read_bytes()[:5] == b"%PDF-":
        raise DocumentError("Le CV doit être un fichier PDF valide.")
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted or not reader.pages:
            raise DocumentError("Le CV PDF est chiffré ou ne contient aucune page.")
    except DocumentError:
        raise
    except Exception as error:
        raise DocumentError("Le CV PDF ne peut pas être lu.") from error


def validate_letter_template(path: Path) -> None:
    """Exige une occurrence unique du marqueur dans tout le DOCX."""
    try:
        document = _load_docx(path)
    except Exception as error:
        raise DocumentError("La lettre doit être un fichier DOCX lisible.") from error
    text = "\n".join(paragraph.text for paragraph in iter_docx_paragraphs(document))
    count = text.count(ROCKY_MARKER)
    if count != 1:
        raise DocumentError(
            "La lettre doit contenir exactement une occurrence de "
            f"{ROCKY_MARKER} ; {count} détectée(s)."
        )


def extract_docx_text(path: Path) -> str:
    """Extrait le texte visible utilisé pour l'analyse et la traduction."""
    document = _load_docx(path)
    return "\n".join(
        paragraph.text.strip()
        for paragraph in iter_docx_paragraphs(document)
        if paragraph.text.strip()
    )


def convert_docx_to_pdf(source: Path, target: Path) -> None:
    """Convertit un DOCX dans un bac LibreOffice isolé puis vérifie le PDF.

    Le profil utilisateur par défaut de LibreOffice est partagé entre les
    sessions Streamlit et rend les conversions concurrentes instables. Chaque
    conversion reçoit donc son propre répertoire de travail et son propre
    profil LibreOffice, puis publie le PDF uniquement après validation.
    """
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise DocumentError(
            "LibreOffice headless est requis pour prévisualiser la lettre DOCX."
        )
    target.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="rocky-libreoffice-") as raw_directory:
        work_dir = Path(raw_directory)
        office_profile = work_dir / "profile"
        office_profile.mkdir()
        # Le nom stable évite les comportements variables de LibreOffice avec
        # les fichiers temporaires cachés tels que ``.letter_upload.docx``.
        work_source = work_dir / "letter.docx"
        shutil.copy2(source, work_source)
        try:
            result = subprocess.run(
                [
                    executable,
                    "--headless",
                    f"-env:UserInstallation={office_profile.as_uri()}",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    str(work_dir),
                    str(work_source),
                ],
                capture_output=True,
                text=True,
                timeout=45,
                check=False,
            )
        except subprocess.TimeoutExpired as error:
            raise DocumentError(
                "La conversion LibreOffice a dépassé le délai autorisé."
            ) from error
        generated = work_dir / "letter.pdf"
        if result.returncode != 0 or not generated.is_file():
            detail = (result.stderr or result.stdout).strip().replace("\n", " ")
            suffix = f" Détail : {detail[:240]}" if detail else ""
            raise DocumentError(
                "LibreOffice n'a pas pu convertir la lettre en PDF." + suffix
            )
        shutil.copy2(generated, target)
    validate_pdf(target)


def _profile_dir(
    settings: Settings, user_id: int, profile_id: int, locale: str
) -> Path:
    """Crée et retourne le répertoire privé d'une langue de profil."""
    directory = settings.user_profiles_dir(user_id) / str(profile_id) / locale
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def save_uploaded_profile_document(
    settings: Settings,
    repository: RockyRepository,
    user_id: int,
    profile_id: int,
    locale: str,
    kind: DocumentKind,
    content: bytes,
) -> Path:
    """Valide puis publie atomiquement le document courant du profil."""
    if locale not in {"fr", "en"}:
        raise ValueError("Document de profil non pris en charge.")
    directory = _profile_dir(settings, user_id, profile_id, locale)
    suffix = ".pdf" if kind == "cv" else ".docx"
    temporary = directory / f".{kind}_upload{suffix}"
    temporary.write_bytes(content)
    preview: Path | None = None
    preview_existed_before = False
    try:
        if kind == "cv":
            validate_pdf(temporary)
        else:
            validate_letter_template(temporary)
        digest = file_sha256(temporary)
        version_tag = digest[:12]
        target = directory / f"{kind}_{version_tag}{suffix}"
        if kind == "letter":
            preview = directory / f"letter_{version_tag}_preview.pdf"
            preview_existed_before = preview.is_file()
            convert_docx_to_pdf(temporary, preview)
        temporary.replace(target)
    except Exception:
        temporary.unlink(missing_ok=True)
        # Une tentative de remplacement ne doit jamais effacer un aperçu
        # validé précédemment pour le même document source.
        if preview is not None and not preview_existed_before:
            preview.unlink(missing_ok=True)
        raise

    stored_source = project_relative(target, settings.project_dir)
    stored_preview = (
        stored_source
        if preview is None
        else project_relative(preview, settings.project_dir)
    )
    repository.save_profile_document(
        profile_id,
        locale,
        kind,
        stored_source,
        digest,
        preview_pdf_path=stored_preview,
        origin="uploaded",
        status="ready",
    )
    if kind == "cv" and locale == "fr":
        repository.save_cv_path(profile_id, stored_source)
    return target


def analyze_profile(
    settings: Settings,
    cv_path: Path,
    letter_path: Path,
    consent_to_llm: bool,
) -> ProfileAnalysis:
    """Combine extraction locale et analyse Mistral, avec repli manuel exploitable."""
    raw_cv, _, _ = extract_pdf_text(cv_path)
    cv_text, _ = repair_spaced_pdf_text(raw_cv)
    letter_text = extract_docx_text(letter_path)
    inferred = infer_profile_skills_from_cv(cv_path)
    local_skills = tuple(skill.name for skill in inferred)
    local_levels = tuple((skill.name, skill.level) for skill in inferred)
    if consent_to_llm:
        llm = RockyLLM(settings)
        if not llm.is_configured:
            raise RockyError("Mistral n'est pas configuré pour analyser les documents.")
        result = llm.analyze_profile_documents(cv_text, letter_text)
        merged_skills = tuple(dict.fromkeys([*result.skills, *local_skills]))
        return ProfileAnalysis(
            full_name=result.full_name,
            email=result.email,
            phone=result.phone,
            summary=result.summary,
            target_job_titles=result.target_job_titles,
            target_domains=result.target_domains,
            skills=merged_skills,
            skill_levels=tuple(dict([*local_levels, *result.skill_levels]).items()),
            career_items=result.career_items,
            project_evidence=result.project_evidence,
            warnings=result.warnings,
        )

    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", cv_text)
    phone_match = re.search(r"(?:\+33|0)[1-9](?:[ .-]?\d{2}){4}", cv_text)
    return ProfileAnalysis(
        email=email_match.group(0) if email_match else "",
        phone=phone_match.group(0) if phone_match else "",
        skills=local_skills,
        skill_levels=local_levels,
        warnings=("Analyse IA désactivée : complète les autres champs manuellement.",),
    )


def _replace_paragraph_text(paragraph, text: str) -> None:
    """Répartit un bloc traduit entre les runs afin de conserver leurs styles."""
    if paragraph.runs:
        original_lengths = [max(1, len(run.text)) for run in paragraph.runs]
        total_weight = sum(original_lengths)
        start = 0
        cumulative = 0
        for index, (run, weight) in enumerate(
            zip(paragraph.runs, original_lengths, strict=True)
        ):
            cumulative += weight
            end = (
                len(text)
                if index == len(paragraph.runs) - 1
                else round(len(text) * cumulative / total_weight)
            )
            run.text = text[start:end]
            start = end
    else:
        paragraph.add_run(text)


def translate_letter_docx(source: Path, target: Path, llm: RockyLLM) -> None:
    """Traduit les blocs DOCX et protège le marqueur Rocky exact."""
    document = _load_docx(source)
    paragraphs = [
        paragraph
        for paragraph in iter_docx_paragraphs(document)
        if paragraph.text.strip()
    ]
    protected = [
        paragraph.text.replace(ROCKY_MARKER, PROTECTED_MARKER)
        for paragraph in paragraphs
    ]
    translated = llm.translate_blocks(protected)
    for paragraph, value in zip(paragraphs, translated, strict=True):
        _replace_paragraph_text(
            paragraph, value.replace(PROTECTED_MARKER, ROCKY_MARKER)
        )
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    validate_letter_template(target)


def fill_letter_template(source: Path, target: Path, rocky_paragraph: str) -> None:
    """Remplace l'unique marqueur en conservant la mise en forme du DOCX source."""
    validate_letter_template(source)
    document = _load_docx(source)
    replaced = False
    for paragraph in iter_docx_paragraphs(document):
        if ROCKY_MARKER not in paragraph.text:
            continue
        _replace_paragraph_text(
            paragraph, paragraph.text.replace(ROCKY_MARKER, rocky_paragraph.strip())
        )
        replaced = True
    if not replaced:
        raise DocumentError("Le paragraphe Rocky n'a pas pu être inséré.")
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))


def render_english_cv(source: Path, target: Path, llm: RockyLLM) -> None:
    """Recompose un CV anglais lisible dans le gabarit sobre Rocky."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph as PdfParagraph
        from reportlab.platypus import SimpleDocTemplate, Spacer
    except ImportError as error:
        raise DocumentError("Installe reportlab pour générer le CV anglais.") from error
    raw_text, _, _ = extract_pdf_text(source)
    repaired, _ = repair_spaced_pdf_text(raw_text)
    blocks = [line.strip() for line in repaired.splitlines() if line.strip()]
    translated: list[str] = []
    for start in range(0, len(blocks), CV_TRANSLATION_BATCH_SIZE):
        translated.extend(
            llm.translate_blocks(blocks[start : start + CV_TRANSLATION_BATCH_SIZE])
        )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "RockyCvBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.4,
        leading=10.2,
        spaceAfter=3,
    )
    heading = ParagraphStyle(
        "RockyCvHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        textColor="#087f96",
        fontSize=11,
        leading=13,
        spaceBefore=7,
        spaceAfter=3,
    )
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=1.25 * cm,
        leftMargin=1.25 * cm,
        topMargin=1.1 * cm,
        bottomMargin=1.1 * cm,
        title="Rocky English CV",
    )
    story = []
    for index, line in enumerate(translated):
        is_heading = index == 0 or (len(line) < 55 and line.upper() == line)
        story.append(PdfParagraph(escape(line), heading if is_heading else body))
        if index == 0:
            story.append(Spacer(1, 4))
    document.build(story)
    validate_pdf(target)


def generate_english_documents(
    settings: Settings,
    repository: RockyRepository,
    user_id: int,
    profile_id: int,
    llm: RockyLLM | None = None,
) -> None:
    """Génère uniquement les actifs anglais absents ou déjà générés.

    Le client injectable sert au smoke test hors ligne ; l'application utilise
    toujours le client Mistral configuré par défaut.
    """
    llm = llm or RockyLLM(settings)
    if not llm.is_configured:
        raise RockyError(
            "Mistral doit être configuré pour générer la version anglaise."
        )
    french = {
        doc.kind: doc for doc in repository.fetch_profile_documents(profile_id, "fr")
    }
    english = {
        doc.kind: doc for doc in repository.fetch_profile_documents(profile_id, "en")
    }
    if "cv" not in french or "letter" not in french:
        raise DocumentError("Ajoute d'abord le CV et la lettre français.")
    source_hash = hashlib.sha256(
        (french["cv"].sha256 + french["letter"].sha256).encode("utf-8")
    ).hexdigest()
    directory = _profile_dir(settings, user_id, profile_id, "en")

    def resolve(value: str) -> Path:
        """Résout un chemin de document historique avant de générer son équivalent anglais."""
        path = Path(value).expanduser()
        return path if path.is_absolute() else settings.project_dir / path

    if "cv" not in english or english["cv"].origin == "generated":
        temporary_cv = directory / ".cv_generated.pdf"
        render_english_cv(resolve(french["cv"].source_path), temporary_cv, llm)
        cv_digest = file_sha256(temporary_cv)
        target_cv = directory / f"cv_{cv_digest[:12]}.pdf"
        temporary_cv.replace(target_cv)
        stored = project_relative(target_cv, settings.project_dir)
        repository.save_profile_document(
            profile_id,
            "en",
            "cv",
            stored,
            cv_digest,
            preview_pdf_path=stored,
            origin="generated",
            status="ready",
            source_hash=source_hash,
        )
    if "letter" not in english or english["letter"].origin == "generated":
        temporary_letter = directory / ".letter_generated.docx"
        translate_letter_docx(
            resolve(french["letter"].source_path), temporary_letter, llm
        )
        letter_digest = file_sha256(temporary_letter)
        target_letter = directory / f"letter_{letter_digest[:12]}.docx"
        target_preview = directory / f"letter_{letter_digest[:12]}_preview.pdf"
        convert_docx_to_pdf(temporary_letter, target_preview)
        temporary_letter.replace(target_letter)
        repository.save_profile_document(
            profile_id,
            "en",
            "letter",
            project_relative(target_letter, settings.project_dir),
            letter_digest,
            preview_pdf_path=project_relative(target_preview, settings.project_dir),
            origin="generated",
            status="ready",
            source_hash=source_hash,
        )
