"""Banc de test ATS V3 multi-parseurs, déterministe et sans profil candidat."""

from __future__ import annotations

import io
import math
import re
import statistics
import subprocess
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from itertools import combinations
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from dashboard.job_analysis import SKILL_ALIASES, analyze_job

from .errors import DocumentError
from .text_utils import normalize_text


ATS_SCREENER_SOURCE = "https://github.com/sunnypatell/ats-screener"


SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "contact": ("contact", "coordonnees", "informations personnelles"),
    "summary": ("profil", "resume", "summary", "a propos", "objectif"),
    "experience": (
        "experience",
        "experiences",
        "experience professionnelle",
        "experiences professionnelles",
        "parcours professionnel",
        "employment history",
        "work experience",
    ),
    "education": (
        "formation",
        "formations",
        "education",
        "diplomes",
        "academic background",
    ),
    "skills": (
        "competences",
        "competences techniques",
        "skills",
        "technical skills",
        "technologies",
        "outils",
    ),
    "projects": ("projets", "projects", "projets personnels"),
    "certifications": (
        "certifications",
        "certification",
        "certificates",
        "licences",
    ),
    "languages": ("langues", "languages", "language proficiency"),
}

DATE_PATTERN = re.compile(
    r"\b(?:"
    r"(?:jan(?:vier|uary)?|f[eé]v(?:rier|ruary)?|mar(?:s|ch)?|avr(?:il)?|apr(?:il)?|"
    r"mai|may|juin|june|juil(?:let)?|july|ao[uû]t|aug(?:ust)?|sep(?:t(?:embre|ember)?)?|"
    r"oct(?:obre|ober)?|nov(?:embre|ember)?|d[eé]c(?:embre|ember)?)\.?\s+)?"
    r"(?:19|20)\d{2}"
    r"(?:\s*[-–—/]\s*(?:(?:19|20)\d{2}|aujourd['’]?hui|present|current))?\b",
    re.IGNORECASE,
)

EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w.-]+\.[a-z]{2,}", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"(?:\+\d{1,3}[ .-]?)?(?:\(?\d{1,3}\)?[ .-]?)?\d(?:[ .-]?\d){7,12}")

LANGUAGE_NAMES = {
    "francais",
    "anglais",
    "espagnol",
    "allemand",
    "italien",
    "portugais",
    "arabe",
    "chinois",
    "japonais",
    "french",
    "english",
    "spanish",
    "german",
    "italian",
    "portuguese",
    "arabic",
    "chinese",
    "japanese",
}

ROLE_WORDS = {
    "analyst",
    "analyste",
    "engineer",
    "ingenieur",
    "scientist",
    "developer",
    "developpeur",
    "consultant",
    "manager",
    "architect",
    "specialist",
    "specialiste",
    "chef",
    "responsable",
    "designer",
}

INSTITUTION_WORDS = {
    "universite",
    "university",
    "ecole",
    "school",
    "institut",
    "institute",
    "bootcamp",
    "college",
    "academie",
    "academy",
}

JOB_STOPWORDS = {
    "avec", "pour", "dans", "vous", "nous", "votre", "notre", "cette",
    "poste", "mission", "missions", "profil", "entreprise", "equipe", "equipes",
    "plus", "ainsi", "etre", "avez", "sera", "sont", "leur", "leurs", "tout",
    "tous", "toutes", "mais", "comme", "chez", "afin", "faire", "fait", "une",
    "des", "les", "sur", "par", "aux", "the", "and", "with", "your", "you",
    "this", "that", "from", "will", "role", "team", "work", "job", "our",
    "are", "have", "has", "into", "about", "they", "their", "who", "what",
}

# Équivalences volontairement courtes et génériques. Elles ne comptent jamais
# comme présence lexicale et ne valent qu'un diagnostic séparé.
SEMANTIC_EQUIVALENCES: dict[str, tuple[str, ...]] = {
    "git": ("gestion de versions", "version control"),
    "jupyter": ("notebook interactif", "interactive notebook"),
    "agile": ("gestion de projet iterative", "iterative project management"),
    "sql": ("requetage relationnel", "relational querying"),
    "data visualisation": ("communication visuelle des donnees",),
    "reporting": ("production de rapports", "report production"),
}


@dataclass(frozen=True)
class StructuredCV:
    name: str
    emails: tuple[str, ...]
    phones: tuple[str, ...]
    professional_title: str
    sections: tuple[str, ...]
    experiences: tuple[str, ...]
    companies: tuple[str, ...]
    dates: tuple[str, ...]
    education: tuple[str, ...]
    institutions: tuple[str, ...]
    skills: tuple[str, ...]
    languages: tuple[str, ...]
    certifications: tuple[str, ...]
    projects: tuple[str, ...]


@dataclass(frozen=True)
class ParserExtraction:
    parser_id: str
    label: str
    engine: str
    license: str
    raw_text: str
    structured: StructuredCV
    quality_score: int
    word_count: int
    character_count: int
    warnings: tuple[str, ...]
    metadata: dict[str, Any]


@dataclass(frozen=True)
class SkillRequirement:
    skill: str
    job_evidence: str
    importance: str


@dataclass(frozen=True)
class SemanticEvidence:
    parser_id: str
    evidence: str


@dataclass(frozen=True)
class SkillComparison:
    skill: str
    importance: str
    job_evidence: str
    exact_parsers: tuple[str, ...]
    variant_parsers: tuple[str, ...]
    semantic_evidence: tuple[SemanticEvidence, ...]
    missing_parsers: tuple[str, ...]


@dataclass(frozen=True)
class BenchmarkResult:
    name: str
    score: int
    interpretation: str
    parsing_component: int
    lexical_component: int
    structure_component: int
    notes: tuple[str, ...]


@dataclass(frozen=True)
class AtsV3Report:
    file_name: str
    file_type: str
    job_title: str
    parser_extractions: tuple[ParserExtraction, ...]
    parsing_robustness: int
    parser_consistency: int
    exact_coverage: int | None
    lexical_coverage: int | None
    mandatory_coverage: int | None
    keyword_coverage: int | None
    semantic_coverage: int | None
    secondary_summary: int | None
    requirements: tuple[SkillRequirement, ...]
    skill_comparisons: tuple[SkillComparison, ...]
    important_keywords: tuple[str, ...]
    benchmark_results: tuple[BenchmarkResult, ...]
    recommendations: tuple[str, ...]
    limits: tuple[str, ...]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _clean_lines(text: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]


def _section_type(line: str) -> str | None:
    cleaned = normalize_text(re.sub(r"[:|_–—-]+$", "", line).strip())
    if not cleaned or len(cleaned.split()) > 6:
        return None
    for section, aliases in SECTION_ALIASES.items():
        if cleaned in aliases:
            return section
    return None


def _sections(lines: list[str]) -> tuple[dict[str, list[str]], list[str]]:
    found: dict[str, list[str]] = {}
    order: list[str] = []
    current = "header"
    found[current] = []
    for line in lines:
        section = _section_type(line)
        if section:
            current = section
            found.setdefault(current, [])
            if section not in order:
                order.append(section)
            continue
        if line:
            found.setdefault(current, []).append(line)
    return found, order


def _extract_name(lines: list[str]) -> str:
    for line in lines[:12]:
        cleaned = re.sub(r"[^A-Za-zÀ-ÿ'’ -]", "", line).strip()
        words = cleaned.split()
        normalized_words = set(normalize_text(cleaned).split())
        if (
            2 <= len(words) <= 4
            and 5 <= len(cleaned) <= 60
            and not normalized_words.intersection(ROLE_WORDS)
            and _section_type(cleaned) is None
            and not EMAIL_PATTERN.search(line)
        ):
            return cleaned
    return ""


def _extract_title(lines: list[str], name: str) -> str:
    for line in lines[:20]:
        normalized_words = set(normalize_text(line).split())
        if (
            normalized_words.intersection(ROLE_WORDS)
            and line != name
            and len(line) <= 100
        ):
            return line
    return ""


def _unique(values: list[str]) -> tuple[str, ...]:
    found: dict[str, str] = {}
    for value in values:
        key = normalize_text(value)
        if key:
            found.setdefault(key, value.strip())
    return tuple(found.values())


def _block_candidates(section_lines: list[str]) -> tuple[str, ...]:
    candidates: list[str] = []
    for index, line in enumerate(section_lines):
        if DATE_PATTERN.search(line):
            start = max(0, index - 1)
            end = min(len(section_lines), index + 2)
            candidates.append(" | ".join(section_lines[start:end]))
    return _unique(candidates)


def _companies(experiences: tuple[str, ...]) -> tuple[str, ...]:
    found: list[str] = []
    for experience in experiences:
        parts = re.split(r"\s[-–—|]\s|\s+chez\s+|\s+at\s+", experience)
        if len(parts) >= 2:
            candidate = DATE_PATTERN.sub("", parts[-1]).strip(" ,-–—|")
            if 2 < len(candidate) < 80:
                found.append(candidate)
    return _unique(found)


def structure_cv(text: str) -> StructuredCV:
    """Structure commune appliquée sans aucune donnée du profil Rocky."""
    lines = _clean_lines(text)
    sections, section_order = _sections(lines)
    header = sections.get("header", lines[:15])
    name = _extract_name(header or lines)
    title = _extract_title(header or lines, name)
    emails = _unique(EMAIL_PATTERN.findall(text))
    phones = _unique(
        match.group(0).strip()
        for match in PHONE_PATTERN.finditer(text)
        if sum(char.isdigit() for char in match.group(0)) >= 9
    )
    dates = _unique(match.group(0) for match in DATE_PATTERN.finditer(text))
    experiences = _block_candidates(sections.get("experience", []))
    education_lines = sections.get("education", [])
    institutions = _unique(
        line
        for line in education_lines
        if set(normalize_text(line).split()).intersection(INSTITUTION_WORDS)
    )
    analysis = analyze_job("", text)
    normalized = normalize_text(text)
    languages = tuple(
        sorted(
            language
            for language in LANGUAGE_NAMES
            if re.search(
                r"(?<![a-z0-9])" + re.escape(language) + r"(?![a-z0-9])",
                normalized,
            )
        )
    )
    return StructuredCV(
        name=name,
        emails=emails,
        phones=phones,
        professional_title=title,
        sections=tuple(section_order),
        experiences=experiences,
        companies=_companies(experiences),
        dates=dates,
        education=_unique(education_lines),
        institutions=institutions,
        skills=tuple(analysis["all_skills"]),
        languages=languages,
        certifications=_unique(sections.get("certifications", [])),
        projects=_unique(sections.get("projects", [])),
    )


def _spaced_character_ratio(text: str) -> float:
    tokens = re.findall(r"\S+", text)
    if not tokens:
        return 1.0
    return sum(len(token) == 1 and token.isalpha() for token in tokens) / len(tokens)


def _garbage_ratio(text: str) -> float:
    if not text:
        return 1.0
    bad = sum(
        not (char.isalnum() or char.isspace() or char in "@.,;:!?%+()/-_'’")
        for char in text
    )
    return bad / len(text)


def _quality(
    text: str, structured: StructuredCV, metadata: dict[str, Any]
) -> tuple[int, tuple[str, ...]]:
    words = re.findall(r"\b\w+\b", text)
    warnings: list[str] = []
    extraction_score = min(100.0, len(text) / 12)
    spacing_ratio = _spaced_character_ratio(text)
    if spacing_ratio > 0.25:
        warnings.append(
            f"{round(spacing_ratio * 100)} % des tokens sont des caractères isolés."
        )
    spacing_score = max(0.0, 100 - spacing_ratio * 140)
    garbage_ratio = _garbage_ratio(text)
    if garbage_ratio > 0.03:
        warnings.append("Caractères inhabituels nombreux dans l’extraction.")
    encoding_score = max(0.0, 100 - garbage_ratio * 800)
    standard_sections = {"experience", "education", "skills"}
    section_score = 100 * len(set(structured.sections) & standard_sections) / 3
    contact_score = 50 * bool(structured.emails) + 50 * bool(structured.phones)
    if not structured.emails:
        warnings.append("Adresse e-mail non détectée.")
    if not structured.phones:
        warnings.append("Téléphone non détecté.")
    if not structured.experiences:
        warnings.append("Aucune expérience structurée détectée.")
    if metadata.get("has_multiple_columns"):
        warnings.append("Disposition multi-colonnes probable.")
    if metadata.get("has_images"):
        warnings.append("Éléments graphiques détectés ; leur texte peut être perdu.")
    score = round(
        0.20 * extraction_score
        + 0.20 * spacing_score
        + 0.15 * encoding_score
        + 0.20 * section_score
        + 0.10 * contact_score
        + 0.15 * min(100, len(structured.dates) * 20)
    )
    if len(words) < 100:
        warnings.append(f"Extraction courte : {len(words)} mots seulement.")
    return max(0, min(100, score)), tuple(warnings)


def _pypdf_extract(data: bytes) -> tuple[str, dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    texts: list[str] = []
    x_positions: list[float] = []
    images = 0
    for page in reader.pages:
        def visitor(text, _cm, tm, _font, _size):
            if str(text).strip():
                x_positions.append(float(tm[4]))

        texts.append(page.extract_text(visitor_text=visitor) or "")
        try:
            images += len(page.images)
        except Exception:
            pass
    clusters = sorted({round(value / 40) * 40 for value in x_positions})
    has_columns = any(
        right - left > 180 for left, right in zip(clusters, clusters[1:])
    ) and len(x_positions) > 30
    return "\n".join(texts).strip(), {
        "page_count": len(reader.pages),
        "has_images": images > 0,
        "image_count": images,
        "has_multiple_columns": has_columns,
    }


def _pdfminer_extract(data: bytes) -> tuple[str, dict[str, Any]]:
    from pdfminer.high_level import extract_text
    from pdfminer.layout import LAParams
    from pdfminer.pdfpage import PDFPage

    stream = io.BytesIO(data)
    page_count = sum(1 for _ in PDFPage.get_pages(stream))
    stream.seek(0)
    text = extract_text(stream, laparams=LAParams())
    return text.strip(), {
        "page_count": page_count,
        "has_images": False,
        "image_count": None,
        "has_multiple_columns": False,
    }


def _pdfium_extract(data: bytes) -> tuple[str, dict[str, Any]]:
    import pypdfium2 as pdfium

    document = pdfium.PdfDocument(data)
    texts: list[str] = []
    for page in document:
        text_page = page.get_textpage()
        texts.append(text_page.get_text_range())
        text_page.close()
        page.close()
    page_count = len(document)
    document.close()
    return "\n".join(texts).strip(), {
        "page_count": page_count,
        "has_images": False,
        "image_count": None,
        "has_multiple_columns": False,
    }


def _docx_paragraph_extract(data: bytes) -> tuple[str, dict[str, Any]]:
    from docx import Document

    document = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return text, {
        "page_count": None,
        "has_images": bool(document.inline_shapes),
        "image_count": len(document.inline_shapes),
        "has_multiple_columns": False,
        "tables_ignored": len(document.tables),
    }


def _docx_xml_extract(data: bytes) -> tuple[str, dict[str, Any]]:
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            root = ElementTree.fromstring(archive.read("word/document.xml"))
            paragraphs = []
            for paragraph in root.iter(namespace + "p"):
                text = "".join(
                    node.text or "" for node in paragraph.iter(namespace + "t")
                ).strip()
                if text:
                    paragraphs.append(text)
            image_count = len(
                [name for name in archive.namelist() if name.startswith("word/media/")]
            )
    except (zipfile.BadZipFile, KeyError, ElementTree.ParseError) as error:
        raise DocumentError("Le fichier DOCX est illisible.") from error
    return "\n".join(paragraphs), {
        "page_count": None,
        "has_images": image_count > 0,
        "image_count": image_count,
        "has_multiple_columns": False,
    }


def _make_extraction(
    parser_id: str,
    label: str,
    engine: str,
    license_name: str,
    text: str,
    metadata: dict[str, Any],
) -> ParserExtraction:
    structured = structure_cv(text)
    quality, warnings = _quality(text, structured, metadata)
    return ParserExtraction(
        parser_id=parser_id,
        label=label,
        engine=engine,
        license=license_name,
        raw_text=text,
        structured=structured,
        quality_score=quality,
        word_count=len(re.findall(r"\b\w+\b", text)),
        character_count=len(text),
        warnings=warnings,
        metadata=metadata,
    )


def extract_with_independent_parsers(
    data: bytes, file_name: str
) -> tuple[ParserExtraction, ...]:
    """Exécute les moteurs tels quels, sans réparation spécifique au CV."""
    suffix = Path(file_name).suffix.lower()
    parsers: list[tuple[str, str, str, str, Any]]
    if suffix == ".pdf":
        parsers = [
            ("pypdf", "pypdf brut", "pypdf", "BSD-3-Clause", _pypdf_extract),
            (
                "pdfminer",
                "pdfminer.six layout",
                "pdfminer.six",
                "MIT",
                _pdfminer_extract,
            ),
            (
                "pdfium",
                "PDFium natif",
                "pypdfium2 / PDFium",
                "Apache-2.0 ou BSD-3-Clause + licences PDFium",
                _pdfium_extract,
            ),
        ]
    elif suffix == ".docx":
        parsers = [
            (
                "docx_paragraphs",
                "Word — paragraphes",
                "python-docx",
                "MIT",
                _docx_paragraph_extract,
            ),
            (
                "docx_ooxml",
                "Word — OOXML brut",
                "Bibliothèque standard XML/ZIP",
                "Python Software Foundation",
                _docx_xml_extract,
            ),
        ]
    else:
        raise DocumentError("ATS V3 accepte uniquement les fichiers PDF ou DOCX.")

    results: list[ParserExtraction] = []
    errors: list[str] = []
    success_count = 0
    for parser_id, label, engine, license_name, parser in parsers:
        try:
            text, metadata = parser(data)
            results.append(
                _make_extraction(
                    parser_id, label, engine, license_name, text, metadata
                )
            )
            success_count += 1
        except Exception as error:
            error_label = f"{label}: {type(error).__name__}"
            errors.append(error_label)
            results.append(
                ParserExtraction(
                    parser_id=parser_id,
                    label=label,
                    engine=engine,
                    license=license_name,
                    raw_text="",
                    structured=structure_cv(""),
                    quality_score=0,
                    word_count=0,
                    character_count=0,
                    warnings=(f"Échec du parseur ({type(error).__name__}).",),
                    metadata={"parser_error": type(error).__name__},
                )
            )
    if success_count < 2:
        raise DocumentError(
            "ATS V3 n’a pas pu exécuter au moins deux parseurs indépendants. "
            + "; ".join(errors)
        )
    return tuple(results)


def _word_set(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-z0-9+#.]{3,}", normalize_text(text))
        if token not in JOB_STOPWORDS
    }


def _jaccard(left: set[str], right: set[str]) -> float:
    union = left | right
    return len(left & right) / len(union) if union else 1.0


def _parser_consistency(extractions: tuple[ParserExtraction, ...]) -> int:
    text_scores = [
        _jaccard(_word_set(left.raw_text), _word_set(right.raw_text))
        for left, right in combinations(extractions, 2)
    ]
    field_scores: list[float] = []
    for left, right in combinations(extractions, 2):
        left_skills = set(left.structured.skills)
        right_skills = set(right.structured.skills)
        skill_score = _jaccard(left_skills, right_skills)
        section_score = _jaccard(
            set(left.structured.sections), set(right.structured.sections)
        )
        contact_score = (
            int(bool(set(left.structured.emails) & set(right.structured.emails)))
            + int(bool(set(left.structured.phones) & set(right.structured.phones)))
        ) / 2
        field_scores.append((skill_score + section_score + contact_score) / 3)
    scores = [*text_scores, *field_scores]
    return round(100 * statistics.mean(scores)) if scores else 100


def _skill_forms(skill: str) -> list[str]:
    return [skill, *SKILL_ALIASES.get(skill, [])]


def _find_phrase(text: str, phrases: list[str]) -> str:
    normalized = normalize_text(text)
    for phrase in sorted(phrases, key=len, reverse=True):
        key = normalize_text(phrase)
        if key and re.search(
            r"(?<![a-z0-9])" + re.escape(key) + r"(?![a-z0-9])", normalized
        ):
            return phrase
    return ""


def _importance(text: str, evidence: str) -> str:
    normalized = normalize_text(text)
    key = normalize_text(evidence)
    position = normalized.find(key)
    if position < 0:
        return "détectée"
    sentence_start = max(
        normalized.rfind(marker, 0, position)
        for marker in (".", "!", "?", "\n", ";")
    )
    sentence_ends = [
        index
        for marker in (".", "!", "?", "\n", ";")
        if (index := normalized.find(marker, position + len(key))) >= 0
    ]
    sentence_end = min(sentence_ends) if sentence_ends else len(normalized)
    context = normalized[sentence_start + 1:sentence_end]
    if any(
        marker in context
        for marker in (
            "indispensable", "obligatoire", "required", "must have", "maitrise",
            "exige", "imperatif", "requis",
        )
    ):
        return "obligatoire"
    if any(
        marker in context
        for marker in ("souhaite", "apprecie", "preferred", "nice to have", "plus")
    ):
        return "souhaitée"
    return "détectée"


def _requirements(job_text: str) -> tuple[SkillRequirement, ...]:
    skills = analyze_job("", job_text)["all_skills"]
    return tuple(
        SkillRequirement(
            skill=skill,
            job_evidence=_find_phrase(job_text, _skill_forms(skill)) or skill,
            importance=_importance(
                job_text, _find_phrase(job_text, _skill_forms(skill)) or skill
            ),
        )
        for skill in skills
    )


def _semantic_evidence(skill: str, text: str) -> str:
    return _find_phrase(text, list(SEMANTIC_EQUIVALENCES.get(normalize_text(skill), ())))


def _compare_skills(
    requirements: tuple[SkillRequirement, ...],
    extractions: tuple[ParserExtraction, ...],
) -> tuple[SkillComparison, ...]:
    comparisons = []
    for requirement in requirements:
        exact: list[str] = []
        variants: list[str] = []
        semantic: list[SemanticEvidence] = []
        missing: list[str] = []
        for extraction in extractions:
            evidence = _find_phrase(
                extraction.raw_text, _skill_forms(requirement.skill)
            )
            if evidence:
                if normalize_text(evidence) == normalize_text(requirement.job_evidence):
                    exact.append(extraction.parser_id)
                else:
                    variants.append(extraction.parser_id)
                continue
            equivalent = _semantic_evidence(requirement.skill, extraction.raw_text)
            if equivalent:
                semantic.append(SemanticEvidence(extraction.parser_id, equivalent))
            else:
                missing.append(extraction.parser_id)
        comparisons.append(
            SkillComparison(
                skill=requirement.skill,
                importance=requirement.importance,
                job_evidence=requirement.job_evidence,
                exact_parsers=tuple(exact),
                variant_parsers=tuple(variants),
                semantic_evidence=tuple(semantic),
                missing_parsers=tuple(missing),
            )
        )
    return tuple(comparisons)


def _coverage(
    comparisons: tuple[SkillComparison, ...],
    parser_count: int,
    *,
    mandatory_only: bool = False,
) -> int | None:
    selected = [
        item
        for item in comparisons
        if not mandatory_only or item.importance == "obligatoire"
    ]
    if not selected:
        return None
    majority = math.ceil(parser_count / 2)
    found = sum(
        len(item.exact_parsers) + len(item.variant_parsers) >= majority
        for item in selected
    )
    return round(100 * found / len(selected))


def _exact_coverage(
    comparisons: tuple[SkillComparison, ...], parser_count: int
) -> int | None:
    """Part des exigences présentes sous la même forme que dans l'annonce."""
    if not comparisons:
        return None
    majority = math.ceil(parser_count / 2)
    found = sum(len(item.exact_parsers) >= majority for item in comparisons)
    return round(100 * found / len(comparisons))


def _semantic_coverage(
    comparisons: tuple[SkillComparison, ...], parser_count: int
) -> int | None:
    missing_lexically = [
        item
        for item in comparisons
        if not item.exact_parsers and not item.variant_parsers
    ]
    if not missing_lexically:
        return None
    majority = math.ceil(parser_count / 2)
    found = sum(len(item.semantic_evidence) >= majority for item in missing_lexically)
    return round(100 * found / len(missing_lexically))


def _important_keywords(job_text: str, limit: int = 18) -> tuple[str, ...]:
    tokens = [
        token
        for token in re.findall(r"[a-z0-9+#.]{4,}", normalize_text(job_text))
        if token not in JOB_STOPWORDS and not token.isdigit()
    ]
    counts: dict[str, int] = {}
    first: dict[str, int] = {}
    for index, token in enumerate(tokens):
        counts[token] = counts.get(token, 0) + 1
        first.setdefault(token, index)
    ranked = sorted(counts, key=lambda token: (-counts[token], first[token]))
    return tuple(ranked[:limit])


def _keyword_coverage(
    keywords: tuple[str, ...], extractions: tuple[ParserExtraction, ...]
) -> int | None:
    if not keywords:
        return None
    majority = math.ceil(len(extractions) / 2)
    found = 0
    for keyword in keywords:
        parsers = sum(
            bool(_find_phrase(extraction.raw_text, [keyword]))
            for extraction in extractions
        )
        found += parsers >= majority
    return round(100 * found / len(keywords))


def _structure_component(extractions: tuple[ParserExtraction, ...]) -> int:
    scores = []
    for extraction in extractions:
        structured = extraction.structured
        signals = (
            bool(structured.emails),
            bool(structured.phones),
            "experience" in structured.sections,
            "education" in structured.sections,
            bool(structured.dates),
            bool(structured.experiences),
        )
        scores.append(100 * sum(signals) / len(signals))
    return round(statistics.mean(scores)) if scores else 0


def _benchmark_results(
    parsing: int,
    exact_coverage: int | None,
    lexical_coverage: int | None,
    semantic_coverage: int | None,
    structure: int,
) -> tuple[BenchmarkResult, ...]:
    exact = float(exact_coverage or 0)
    lexical = float(lexical_coverage or 0)
    semantic = float(semantic_coverage or 0)
    profiles = (
        (
            "Workday-like",
            0.40 * parsing + 0.40 * exact + 0.20 * structure,
            parsing,
            round(exact),
            structure,
            ("Parsing strict", "Priorité aux termes littéraux et sections standard"),
        ),
        (
            "Taleo / Oracle-like",
            0.25 * parsing + 0.55 * exact + 0.20 * structure,
            parsing,
            round(exact),
            structure,
            ("Filtrage littéral renforcé", "Faible tolérance aux variantes"),
        ),
        (
            "iCIMS-like",
            0.25 * parsing + 0.50 * max(lexical, semantic) + 0.25 * structure,
            parsing,
            round(max(lexical, semantic)),
            structure,
            ("Variantes et équivalences séparées", "Structure toujours nécessaire"),
        ),
        (
            "Greenhouse-like",
            0.25 * parsing + 0.30 * lexical + 0.45 * structure,
            parsing,
            round(lexical),
            structure,
            ("Indicateur de visibilité, pas prédiction de classement", "Lecture humaine importante"),
        ),
        (
            "Lever-like",
            0.20 * parsing + 0.40 * lexical + 0.40 * structure,
            parsing,
            round(lexical),
            structure,
            ("Recherche textuelle tolérante", "Contexte d’expérience valorisé"),
        ),
        (
            "SuccessFactors-like",
            0.35 * parsing + 0.30 * exact + 0.35 * structure,
            parsing,
            round(exact),
            structure,
            ("Cartographie structurée des champs", "Dates et rubriques déterminantes"),
        ),
    )
    results = []
    for name, score, parse_score, lexical_score, structure_score, notes in profiles:
        rounded = round(score)
        interpretation = (
            "robuste" if rounded >= 75 else "intermédiaire" if rounded >= 55 else "fragile"
        )
        results.append(
            BenchmarkResult(
                name=name,
                score=rounded,
                interpretation=interpretation,
                parsing_component=parse_score,
                lexical_component=lexical_score,
                structure_component=structure_score,
                notes=notes,
            )
        )
    return tuple(results)


def _recommendations(
    extractions: tuple[ParserExtraction, ...],
    comparisons: tuple[SkillComparison, ...],
    consistency: int,
) -> tuple[str, ...]:
    parser_count = len(extractions)
    recommendations: list[str] = []
    for item in comparisons:
        lexical_count = len(item.exact_parsers) + len(item.variant_parsers)
        if lexical_count == 0:
            recommendations.append(
                f"{item.skill} est demandé dans l’annonce mais absent du texte "
                f"extrait par les {parser_count} parseurs."
            )
        elif lexical_count < parser_count:
            recommendations.append(
                f"{item.skill} n’est reconnu que par {lexical_count} parseur(s) "
                f"sur {parser_count} : vérifier sa position et son écriture dans le CV."
            )
    email_counts = sum(bool(item.structured.emails) for item in extractions)
    phone_counts = sum(bool(item.structured.phones) for item in extractions)
    if email_counts < parser_count:
        recommendations.append(
            f"L’e-mail n’est lu que par {email_counts} parseur(s) sur {parser_count}."
        )
    if phone_counts < parser_count:
        recommendations.append(
            f"Le téléphone n’est lu que par {phone_counts} parseur(s) sur {parser_count}."
        )
    experience_counts = [len(item.structured.experiences) for item in extractions]
    if len(set(experience_counts)) > 1:
        detail = ", ".join(
            f"{item.label}: {len(item.structured.experiences)}"
            for item in extractions
        )
        recommendations.append(
            "Le nombre d’expériences reconnues diverge entre parseurs (" + detail + ")."
        )
    date_counts = [len(item.structured.dates) for item in extractions]
    if len(set(date_counts)) > 1:
        recommendations.append(
            "Les dates ne sont pas extraites de façon cohérente : "
            + ", ".join(
                f"{item.label}: {len(item.structured.dates)}"
                for item in extractions
            )
            + "."
        )
    if consistency < 70:
        recommendations.append(
            f"La cohérence inter-parseurs est de {consistency} % : comparer les "
            "vues brutes pour repérer un ordre de lecture ou une colonne problématique."
        )
    for extraction in extractions:
        if _spaced_character_ratio(extraction.raw_text) > 0.25:
            recommendations.append(
                f"{extraction.label} restitue beaucoup de caractères isolés ; le "
                "tracking typographique du PDF présente un risque réel."
            )
    if not recommendations:
        recommendations.append(
            "Aucune divergence majeure mesurable ; conserver néanmoins une version "
            "DOCX ou PDF mono-colonne pour les portails sensibles."
        )
    return tuple(dict.fromkeys(recommendations))


def analyze_ats_v3(
    cv_data: bytes,
    file_name: str,
    job_description: str,
    job_title: str = "",
) -> AtsV3Report:
    """Analyse le fichier réel sans compléter depuis le profil Rocky ni un LLM."""
    if not cv_data:
        raise DocumentError("Le fichier CV est vide.")
    if len(job_description.strip()) < 80:
        raise DocumentError(
            "La description d’annonce est trop courte pour un test V3 robuste."
        )
    extractions = extract_with_independent_parsers(cv_data, file_name)
    consistency = _parser_consistency(extractions)
    average_quality = round(
        statistics.mean(item.quality_score for item in extractions)
    )
    parsing = round(0.55 * average_quality + 0.45 * consistency)
    requirements = _requirements(job_description)
    comparisons = _compare_skills(requirements, extractions)
    exact = _exact_coverage(comparisons, len(extractions))
    lexical = _coverage(comparisons, len(extractions))
    mandatory = _coverage(
        comparisons, len(extractions), mandatory_only=True
    )
    semantic = _semantic_coverage(comparisons, len(extractions))
    keywords = _important_keywords(job_description)
    keyword_coverage = _keyword_coverage(keywords, extractions)
    structure = _structure_component(extractions)
    benchmark = _benchmark_results(
        parsing, exact, lexical, semantic, structure
    )
    if lexical is None:
        secondary = None
    else:
        secondary = round(
            0.45 * parsing
            + 0.40 * lexical
            + 0.15 * (keyword_coverage or 0)
        )
    return AtsV3Report(
        file_name=file_name,
        file_type=Path(file_name).suffix.lower().lstrip("."),
        job_title=job_title,
        parser_extractions=extractions,
        parsing_robustness=parsing,
        parser_consistency=consistency,
        exact_coverage=exact,
        lexical_coverage=lexical,
        mandatory_coverage=mandatory,
        keyword_coverage=keyword_coverage,
        semantic_coverage=semantic,
        secondary_summary=secondary,
        requirements=requirements,
        skill_comparisons=comparisons,
        important_keywords=keywords,
        benchmark_results=benchmark,
        recommendations=_recommendations(extractions, comparisons, consistency),
        limits=(
            "Les benchmarks sont des comportements inspirés, pas les algorithmes propriétaires réels.",
            "La structuration commune repose sur des règles génériques français/anglais.",
            "Aucun OCR n’est appliqué : un CV scanné doit échouer visiblement au lieu d’être réparé en silence.",
            "La couche sémantique est une petite table explicite et ne modifie jamais la couverture lexicale.",
        ),
    )


def render_pdf_first_page(cv_data: bytes, scale: float = 1.4) -> bytes:
    """Rend la première page via PDFium pour confronter visuel et extraction."""
    import pypdfium2 as pdfium

    document = pdfium.PdfDocument(cv_data)
    if len(document) == 0:
        raise DocumentError("Le PDF ne contient aucune page.")
    page = document[0]
    image = page.render(scale=scale).to_pil()
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    page.close()
    document.close()
    return buffer.getvalue()


def poppler_diagnostic(cv_data: bytes) -> str | None:
    """Diagnostic externe optionnel ; absent de la note si Poppler manque."""
    try:
        with tempfile.TemporaryDirectory(prefix="rocky_ats_v3_") as directory:
            source = Path(directory) / "cv.pdf"
            output = Path(directory) / "cv.txt"
            source.write_bytes(cv_data)
            subprocess.run(
                ["pdftotext", "-layout", str(source), str(output)],
                check=True,
                capture_output=True,
                timeout=15,
            )
            return output.read_text(encoding="utf-8", errors="replace")
    except (FileNotFoundError, subprocess.SubprocessError, OSError):
        return None
