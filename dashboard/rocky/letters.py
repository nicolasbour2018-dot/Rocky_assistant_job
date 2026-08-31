"""Prévisualisation et génération des lettres de motivation."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from html import escape
from pathlib import Path
from string import Formatter

from .config import Settings
from .errors import DocumentError


@dataclass(frozen=True)
class LetterVariables:
    """Variables factuelles injectées dans le modèle de lettre d'une candidature."""
    job_title: str
    company_name: str
    company_paragraph: str
    recipient: str = "À l’attention du Service des Ressources Humaines"
    company_address: str = ""
    city: str = ""
    letter_date: date = field(default_factory=date.today)
    sender_name: str = ""
    sender_address: str = ""
    sender_phone: str = ""
    sender_email: str = ""
    locale: str = "fr"

    def template_values(self) -> dict[str, str]:
        """Expose les variables contrôlées injectables dans le modèle de lettre."""
        months = [
            "janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre",
        ]
        french_date = (
            f"{self.letter_date.day} {months[self.letter_date.month - 1]} "
            f"{self.letter_date.year}"
        )
        displayed_date = (
            self.letter_date.strftime("%B %d, %Y")
            if self.locale == "en"
            else french_date
        )
        return {
            "job_title": self.job_title.strip(),
            "company_name": self.company_name.strip(),
            "company_paragraph": self.company_paragraph.strip(),
            "recipient": self.recipient.strip(),
            "company_address": self.company_address.strip(),
            "city_and_date": (
                f"{self.city.strip()}, {displayed_date}"
                if self.locale == "en"
                else f"{self.city.strip()}, le {displayed_date}"
            ),
            "sender_name": self.sender_name,
            "sender_address": self.sender_address,
            "sender_phone": self.sender_phone,
            "sender_email": self.sender_email,
        }


@dataclass(frozen=True)
class LetterSections:
    """Blocs éditables d'une lettre tout en conservant sa mise en page."""

    sender: str
    recipient: str
    city_and_date: str
    subject: str
    salutation: str
    body: tuple[str, ...]


def _template(settings: Settings) -> str:
    """Charge le modèle de lettre validé, sans en modifier le texte métier."""
    return (
        settings.project_dir / "templates" / "lettre_motivation.txt"
    ).read_text(encoding="utf-8")


def render_letter(settings: Settings, variables: LetterVariables) -> str:
    """Remplace uniquement les variables autorisées du modèle."""
    if variables.locale == "en":
        return render_letter_from_body(
            variables,
            (
                f"I am applying for the {variables.job_title.strip()} position "
                f"at {variables.company_name.strip()}.",
                variables.company_paragraph.strip(),
                "My attached resume presents the experience, skills and projects "
                "supporting this application.",
                "I would welcome the opportunity to discuss the role and my application.",
            ),
        )
    template = _template(settings)
    allowed = set(variables.template_values())
    requested = {
        name
        for _, name, _, _ in Formatter().parse(template)
        if name is not None
    }
    if requested != allowed:
        raise DocumentError(
            "Les variables du modèle de lettre ne correspondent pas au code."
        )
    values = variables.template_values()
    if not values["job_title"] or not values["company_name"]:
        raise DocumentError("Le poste et l'entreprise sont obligatoires.")
    if not values["company_paragraph"]:
        raise DocumentError("Le paragraphe entreprise est obligatoire.")
    result = template.format(**values).strip() + "\n"
    if "{" in result or "}" in result:
        raise DocumentError("Une variable non remplacée reste dans la lettre.")
    return result


def render_letter_from_body(
    variables: LetterVariables, body_paragraphs: tuple[str, ...]
) -> str:
    """Assemble un corps ciblé dans la structure PDF validée de Rocky."""
    values = variables.template_values()
    if not 4 <= len(body_paragraphs) <= 6:
        raise DocumentError("La lettre ciblée doit contenir 4 à 6 paragraphes.")
    recipient_lines = [values["recipient"], values["company_name"]]
    if values["company_address"]:
        recipient_lines.append(values["company_address"])
    english = variables.locale == "en"
    blocks = [
        "\n".join(
            [
                values["sender_name"],
                values["sender_address"],
                values["sender_phone"],
                values["sender_email"],
            ]
        ),
        "\n".join(recipient_lines),
        values["city_and_date"],
        (
            f"Subject: Application for the {values['job_title']} position"
            if english
            else f"Objet : Candidature pour le poste de {values['job_title']}"
        ),
        "Dear Hiring Team," if english else "Madame, Monsieur,",
        *[" ".join(paragraph.split()) for paragraph in body_paragraphs],
        (
            "Yours sincerely,"
            if english
            else "Je vous prie d’agréer, Madame, Monsieur, l’expression de mes salutations distinguées."
        ),
        values["sender_name"],
    ]
    return "\n\n".join(blocks).strip() + "\n"


def parse_letter_sections(letter_text: str) -> LetterSections:
    """Découpe une lettre modifiée en blocs destinés au DOCX et au PDF.

    L'utilisateur peut modifier librement le contenu de chaque bloc. Les lignes
    blanches délimitent l'expéditeur, le destinataire, la date, l'objet, la
    salutation puis les paragraphes du corps.
    """
    blocks = [
        block.strip()
        for block in re.split(r"\n\s*\n", letter_text.strip())
        if block.strip()
    ]
    if len(blocks) < 6:
        raise DocumentError(
            "La lettre doit conserver une ligne vide entre l'expéditeur, le "
            "destinataire, la date, l'objet, la salutation et le corps."
        )
    return LetterSections(
        sender=blocks[0],
        recipient=blocks[1],
        city_and_date=blocks[2],
        subject=blocks[3],
        salutation=blocks[4],
        body=tuple(blocks[5:]),
    )


def render_letter_preview_html(
    variables: LetterVariables, letter_text: str
) -> str:
    """Construit une prévisualisation fidèle sans injecter de contenu HTML.

    Toutes les valeurs saisies sont échappées avant leur affichage. Le HTML ne
    sert ici qu'à reproduire les alignements du DOCX et du PDF dans Streamlit.
    """
    sections = parse_letter_sections(letter_text)

    def lines(value: str) -> str:
        """Échappe et conserve les sauts de ligne pour l'aperçu HTML de la lettre."""
        return "<br>".join(
            escape(line) for line in value.splitlines() if line.strip()
        )

    sender = lines(sections.sender)
    recipient = lines(sections.recipient)
    sender_name = sections.sender.splitlines()[0].strip()
    body_paragraphs = []
    for text in sections.body:
        alignment = "left" if text == sender_name else "justify"
        body_paragraphs.append(
            f'<p style="text-align:{alignment}; margin:0 0 0.65rem 0;">'
            f"{escape(text)}</p>"
        )

    return f"""
    <div style="
        max-width: none;
        margin: 0 auto;
        padding: 2rem 2.4rem;
        border: 1px solid #d8d8d8;
        background: white;
        color: #111;
        font-family: Arial, sans-serif;
        font-size: 0.95rem;
        line-height: 1.28;
    ">
        <div style="text-align:left; margin-bottom:1rem;">{sender}</div>
        <div style="text-align:right; margin-bottom:0.8rem;">{recipient}</div>
        <div style="text-align:right; margin-bottom:1rem;">
            {escape(sections.city_and_date)}
        </div>
        <div style="text-align:left; font-weight:bold; margin-bottom:1rem;">
            {escape(sections.subject)}
        </div>
        <p style="text-align:left; margin:0 0 0.65rem 0;">
            {escape(sections.salutation)}
        </p>
        {''.join(body_paragraphs)}
    </div>
    """


def _add_docx_paragraph(document, text: str, alignment=None, bold=False):
    """Ajoute un paragraphe DOCX homogène pendant le rendu de la lettre finale."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = alignment or WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4.5)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(10)
    return paragraph


def create_docx(path: Path, variables: LetterVariables, letter_text: str) -> None:
    """Crée un DOCX A4 sobre, proche de la lettre de référence."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as error:
        raise DocumentError(
            "Installe python-docx pour générer le fichier éditable."
        ) from error

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.0

    values = variables.template_values()
    sections = parse_letter_sections(letter_text)
    sender_name = sections.sender.splitlines()[0].strip()
    _add_docx_paragraph(document, sections.sender)
    # Le bloc destinataire contient notamment le nom de l'entreprise et reste
    # aligné à droite, comme dans une lettre professionnelle classique.
    _add_docx_paragraph(
        document, sections.recipient, alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )
    _add_docx_paragraph(
        document,
        sections.city_and_date,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _add_docx_paragraph(
        document,
        sections.subject,
        bold=True,
    )

    _add_docx_paragraph(document, sections.salutation)
    for text in sections.body:
        alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if text == sender_name
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        _add_docx_paragraph(
            document, text, alignment=alignment
        )

    document.core_properties.title = (
        f"Candidature - {values['job_title']} - {values['company_name']}"
    )
    document.core_properties.author = values["sender_name"]
    document.save(path)


def create_pdf(path: Path, variables: LetterVariables, letter_text: str) -> None:
    """Crée le PDF final indépendamment du logiciel Word."""
    try:
        from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as error:
        raise DocumentError(
            "Installe reportlab pour générer le fichier PDF."
        ) from error
    from xml.sax.saxutils import escape

    values = variables.template_values()
    sections = parse_letter_sections(letter_text)
    sender_name = sections.sender.splitlines()[0].strip()
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.7 * cm,
        rightMargin=1.7 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.2 * cm,
        title=f"Candidature - {values['job_title']}",
        author=values["sender_name"],
    )
    left_style = ParagraphStyle(
        "LetterLeft",
        fontName="Helvetica",
        fontSize=10,
        leading=11.7,
        spaceAfter=5,
        alignment=TA_LEFT,
    )
    body_style = ParagraphStyle(
        "LetterBody",
        parent=left_style,
        alignment=TA_JUSTIFY,
    )
    right_style = ParagraphStyle(
        "LetterRight",
        parent=left_style,
        alignment=TA_RIGHT,
        spaceAfter=7,
    )
    subject_style = ParagraphStyle(
        "LetterSubject",
        parent=left_style,
        fontName="Helvetica-Bold",
        spaceAfter=8,
    )

    sender = "<br/>".join(
        escape(line)
        for line in sections.sender.splitlines()
        if line.strip()
    )
    recipient = "<br/>".join(
        escape(line)
        for line in sections.recipient.splitlines()
        if line.strip()
    )
    story = [
        Paragraph(sender, left_style),
        Spacer(1, 3),
        Paragraph(recipient, right_style),
        Paragraph(escape(sections.city_and_date), right_style),
        Paragraph(escape(sections.subject), subject_style),
        Paragraph(escape(sections.salutation), left_style),
    ]
    for text in sections.body:
        style = left_style if text == sender_name else body_style
        story.append(Paragraph(escape(text), style))
    document.build(story)


def save_profile_cv(
    settings: Settings, profile_id: int, content: bytes
) -> Path:
    """Enregistre le CV du profil dans un emplacement privé et stable."""
    if not content.startswith(b"%PDF"):
        raise DocumentError("Le CV doit être un fichier PDF valide.")
    profile_dir = settings.profiles_dir / str(profile_id)
    profile_dir.mkdir(parents=True, exist_ok=True)
    target = profile_dir / "cv.pdf"
    target.write_bytes(content)
    return target
